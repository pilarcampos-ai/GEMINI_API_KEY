import fitz  # PyMuPDF
from docx import Document
import google.generativeai as genai
import os
import glob

# Configurar la Inteligencia Artificial
genai.configure(api_key=os.getenv("API_KEY"))
model = genai.GenerativeModel('gemini-1.5-flash')

def traducir_texto(texto):
    if not texto.strip(): return ""
    # Instrucciones precisas para la IA
    prompt = f"Traduce fielmente este manual técnico al español. Mantén estructura, listas y términos técnicos. No resumas:\n\n{texto}"
    try:
        response = model.generate_content(prompt)
        return response.text
    except:
        return texto # Si falla, deja el original para no perder datos

# Revisar si hay archivos PDF en la carpeta 'entrada'
archivos_pdf = glob.glob("entrada/*.pdf")

if not archivos_pdf:
    print("No se encontraron archivos PDF en la carpeta 'entrada'.")
else:
    for ruta_pdf in archivos_pdf:
        nombre_base = os.path.basename(ruta_pdf).replace(".pdf", "")
        print(f"Traduciendo: {nombre_base}...")
        
        pdf = fitz.open(ruta_pdf)
        doc_word = Document()
        doc_word.add_heading(f'Manual Traducido: {nombre_base}', 0)

        for pagina in pdf:
            bloques = pagina.get_text("blocks")
            for b in bloques:
                texto_original = b[4]
                if len(texto_original.strip()) > 1:
                    texto_es = traducir_texto(texto_original)
                    doc_word.add_paragraph(texto_es)
        
        # Guardar el resultado en la carpeta 'salida'
        doc_word.save(f"salida/TRADUCIDO_{nombre_base}.docx")
        print(f"¡Listo! Guardado en la carpeta 'salida'.")
